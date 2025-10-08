from docx import Document
from docx.shared import Pt


def add_heading(document: Document, text: str, level: int = 2) -> None:
    document.add_heading(text, level=level)


def add_paragraph(document: Document, text: str) -> None:
    document.add_paragraph(text)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_note_section(document: Document, title: str, notes: list[str]) -> None:
    p = document.add_paragraph()
    run = p.add_run(title)
    run.italic = True
    run.font.size = Pt(11)
    add_bullets(document, notes)


def build_document() -> Document:
    document = Document()

    # Titel
    document.add_heading("HiveMQ – Unternehmen, Produkt, Markt & Karriere", 0)
    intro = document.add_paragraph()
    intro_run = intro.add_run(
        "Zielgruppe: Studierende der Wirtschaftsinformatik • Ziel: HiveMQ vorstellen – Nutzenversprechen, Markt, Produkt, Geschäftsmodell, Karriereperspektiven • Format: 3 Sprecher:innen × 10 Minuten + kurze Q&A"
    )
    intro_run.font.size = Pt(11)

    # Gesamtagenda
    add_heading(document, "Gesamtagenda (30 Minuten)", level=1)
    add_bullets(
        document,
        [
            "Sprecher:in 1 (10 Min): Unternehmen & Markt",
            "Sprecher:in 2 (10 Min): Produkt & Technologie (MQTT, Plattform, Value)",
            "Sprecher:in 3 (10 Min): Geschäftsmodell, Go‑to‑Market, Wettbewerb, Karriere",
        ],
    )

    # Sprecher:in 1
    add_heading(document, "Sprecher:in 1 – Unternehmen & Markt (10 Min)", level=1)

    # Folie 1
    add_heading(document, "Folie 1 (1 Min): Titel & Hook")
    add_bullets(
        document,
        [
            "Key Message: Zuverlässige IoT‑Datenübertragung in großem Maßstab – das ist HiveMQ.",
            "Visual: Bild einer vernetzten Fabrik oder eines Connected Cars",
        ],
    )

    # Folie 2
    add_heading(document, "Folie 2 (3 Min): Kurzprofil Unternehmen")
    add_bullets(
        document,
        [
            "Was: MQTT‑Plattform/Message Broker für IoT",
            "Herkunft: Tech‑Unternehmen aus Deutschland, international tätig",
            "Kundenbranchen: Automotive, Industrie 4.0, Logistik, Energie, Smart Devices",
            "Wertversprechen: Zuverlässigkeit, Skalierbarkeit, Interoperabilität, Sicherheit, Observability",
        ],
    )

    # Folie 3
    add_heading(document, "Folie 3 (3 Min): Markt & Problem")
    add_bullets(
        document,
        [
            "IoT‑Trend: Milliarden Geräte, heterogene Netze, instabile Verbindungen",
            "Problem: Sichere, performante, standardbasierte Kommunikation Gerät ↔ Cloud ↔ Systeme",
            "Lösung: MQTT (OASIS‑Standard), leichtgewichtig, QoS, Retained Messages, Sessions",
        ],
    )

    # Folie 4
    add_heading(document, "Folie 4 (3 Min): Positionierung von HiveMQ")
    add_bullets(
        document,
        [
            "Rolle: Enterprise‑MQTT‑Plattform (on‑prem, Cloud, Hybrid)",
            "Differenzierung: Enterprise‑Features, Cluster‑Skalierung, Data Quality (Data Hub), Ökosystem/Extensions",
            "Outcomes: Schnellere Time‑to‑Value, geringeres Integrationsrisiko, Betriebssicherheit",
            "Übergabe zu Sprecher:in 2: Vom Markt/Problem zur Technologie & Plattform",
        ],
    )

    add_note_section(
        document,
        "Sprechernotizen (kurz):",
        [
            "Business‑Outcome (z. B. OEE, Kundenerlebnis) mit technischer Basis (MQTT) verbinden",
            "Standard‑Compliance und Vendor‑Neutralität betonen",
        ],
    )

    # Sprecher:in 2
    add_heading(document, "Sprecher:in 2 – Produkt & Technologie (10 Min)", level=1)

    # Folie 5
    add_heading(document, "Folie 5 (2 Min): MQTT Basics in 60 Sekunden")
    add_bullets(
        document,
        [
            "Publish/Subscribe‑Modell und Topics",
            "QoS 0/1/2",
            "Retained Messages",
            "Last Will",
        ],
    )

    # Folie 6
    add_heading(document, "Folie 6 (3 Min): HiveMQ Plattformüberblick")
    add_bullets(
        document,
        [
            "Komponenten: Broker/Cluster, HiveMQ Cloud (Managed), Data Hub (Validierung/Transformation), Extensions SDK, Observability",
            "Deployment: Kubernetes/VM, Multi‑Cloud, Edge‑Integration",
        ],
    )

    # Folie 7
    add_heading(document, "Folie 7 (3 Min): Enterprise‑Eigenschaften")
    add_bullets(
        document,
        [
            "Zuverlässigkeit: Hohe Verfügbarkeit, Stateful Sessions, Backpressure",
            "Sicherheit: TLS, AuthN/AuthZ, Zertifikate, Richtlinien",
            "Operability: Metrics, Tracing, Integrationen (Data Lakes, Analytics, iPaaS)",
        ],
    )

    # Folie 8
    add_heading(document, "Folie 8 (2 Min): Kurzer Use Case Flow")
    add_bullets(
        document,
        [
            "Beispiel: Connected Factory – Sensor → MQTT → HiveMQ → Analytics/ERP",
            "Value: Echtzeit‑Monitoring, geringere Ausfallzeiten, Datenqualität",
            "Optional: Mini‑Demo (MQTT‑Client Publish/Subscribe gegen Test‑Broker; Topic‑Flow & QoS zeigen)",
            "Übergabe zu Sprecher:in 3: Von Technik zu Geschäftsmodell, GTM & Karriere",
        ],
    )

    add_note_section(
        document,
        "Sprechernotizen:",
        [
            "Nicht zu tief in Protokolldetails – Fokus auf Betrieb: Latenz, Zuverlässigkeit, Operability",
        ],
    )

    # Sprecher:in 3
    add_heading(document, "Sprecher:in 3 – Geschäftsmodell, GTM, Wettbewerb, Karriere (10 Min)", level=1)

    # Folie 9
    add_heading(document, "Folie 9 (3 Min): Geschäftsmodell & Pricing‑Hebel")
    add_bullets(
        document,
        [
            "Erlösmodell: Subscription/Enterprise‑Lizenz, Managed Service (Cloud)",
            "Werttreiber: Verbindungsanzahl/Throughput, Cluster‑Größe, SLA, Add‑ons (Data Hub)",
            "TCO: Build‑vs‑Buy – geringere Betriebs‑ und Fehlerrisiken",
        ],
    )

    # Folie 10
    add_heading(document, "Folie 10 (3 Min): Go‑to‑Market & Wettbewerb")
    add_bullets(
        document,
        [
            "GTM: Direktvertrieb Enterprise, Partner/Integratoren, Cloud‑Angebot",
            "Wettbewerb: Open‑Source‑Broker (z. B. Mosquitto, VerneMQ), kommerzielle (z. B. EMQX), Hyperscaler‑Dienste (z. B. AWS/Azure IoT)",
            "Differenzierung: Enterprise‑Reliability, Operability, Data Quality, Support",
        ],
    )

    # Folie 11
    add_heading(document, "Folie 11 (2 Min): Karrierepfade für Wirtschaftsinformatik")
    add_bullets(
        document,
        [
            "Rollen: Product Management, Solutions/Pre‑Sales, Customer Success, Data/Cloud Engineering, Partner Management",
            "Skill‑Fit: Schnittstelle Business‑Tech, Architekturen, Datenintegration, Security/Governance",
        ],
    )

    # Folie 12
    add_heading(document, "Folie 12 (2 Min): Takeaways & Q&A")
    add_bullets(
        document,
        [
            "HiveMQ = Brückentechnologie zwischen Geräten und Business‑Systemen",
            "Standardbasiert, skalierbar, betriebssicher, beschleunigt IoT‑Value",
            "Kurze Q&A",
        ],
    )

    add_note_section(
        document,
        "Sprechernotizen:",
        [
            "Klarstellen, wie WiInf Mehrwert schafft: Business‑Ziele in Datenflüsse/Policies/SLAs übersetzen",
            "Storyline betonen: Warum (Markt/Problem) → Wie (MQTT & Plattform) → Wert (Business Outcomes & TCO)",
            "Klare Übergaben: Markt → Technik → Business/Karriere",
        ],
    )

    # Optional: Weiterführende Ressourcen (Handout)
    add_heading(document, "Weiterführende Ressourcen (Handout)", level=1)
    add_bullets(
        document,
        [
            "MQTT Standard (OASIS): Grundlagen und QoS‑Konzepte",
            "HiveMQ Website: Produkt & Ressourcen",
            "Best Practices: Architektur & Operations für IoT‑Skalierung",
        ],
    )

    return document


def main() -> None:
    document = build_document()
    output_path = "/workspace/HiveMQ_Praesentation.docx"
    document.save(output_path)


if __name__ == "__main__":
    main()

