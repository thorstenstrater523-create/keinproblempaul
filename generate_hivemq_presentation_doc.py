from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def add_title_page(document: Document) -> None:
    document.add_heading("HiveMQ – Unternehmensvorstellung für Wirtschaftsinformatik-Studierende", 0)
    subtitle = document.add_paragraph()
    run = subtitle.add_run("3 Sprecher:innen × 10 Minuten | " + date.today().strftime("%d.%m.%Y"))
    run.italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Number")
        p.add_run(item)


def add_notes(document: Document, notes: list[str]) -> None:
    p = document.add_paragraph()
    title = p.add_run("Sprechernotizen: ")
    title.bold = True
    for idx, note in enumerate(notes, start=1):
        note_p = document.add_paragraph(style="List Bullet")
        run = note_p.add_run(note)
        run.italic = True


def build_document() -> Document:
    doc = Document()

    # Title page
    add_title_page(doc)

    # Agenda
    add_heading(doc, "Agenda", level=1)
    add_bullets(
        doc,
        [
            "Sprecher:in 1 – Unternehmen & Markt (10 Min)",
            "Sprecher:in 2 – Produkt & Technologie (10 Min)",
            "Sprecher:in 3 – Geschäftsmodell, Go-to-Market, Karriere (10 Min)",
        ],
    )

    doc.add_page_break()

    # Sprecher:in 1 – Unternehmen & Markt
    add_heading(doc, "Sprecher:in 1 – Unternehmen & Markt", level=1)

    add_heading(doc, "Folie 1 – Titel & Hook", level=2)
    add_bullets(
        doc,
        [
            "Zuverlässige IoT-Datenübertragung in großem Maßstab – das ist HiveMQ.",
            "Visual: Vernetzte Fabrik / Connected Car als Einstimmung.",
        ],
    )

    add_heading(doc, "Folie 2 – Kurzprofil Unternehmen", level=2)
    add_bullets(
        doc,
        [
            "Was: MQTT-Plattform/Message Broker für IoT (Enterprise)",
            "Herkunft: Tech-Unternehmen aus Deutschland, international tätig",
            "Kundenbranchen: Automotive, Industrie 4.0, Logistik, Energie, Smart Devices",
            "Wertversprechen: Zuverlässigkeit, Skalierbarkeit, Interoperabilität, Sicherheit, Observability",
        ],
    )

    add_heading(doc, "Folie 3 – Markt & Problem", level=2)
    add_bullets(
        doc,
        [
            "IoT-Trend: Milliarden Geräte, heterogene Netze, schwankende Verbindungen",
            "Problem: Sichere, performante, standardbasierte Kommunikation Gerät ↔ Cloud ↔ Systeme",
            "Lösung: MQTT (OASIS-Standard) – leichtgewichtig, QoS, Retained Messages, Sessions",
        ],
    )

    add_heading(doc, "Folie 4 – Positionierung von HiveMQ", level=2)
    add_bullets(
        doc,
        [
            "Rolle: Enterprise‑MQTT‑Plattform (on‑prem, Cloud, Hybrid)",
            "Differenzierung: Enterprise‑Features, Cluster‑Skalierung, Data Quality (Data Hub), Extensions",
            "Outcomes: Time‑to‑Value, geringeres Integrationsrisiko, Betriebssicherheit",
        ],
    )
    add_notes(
        doc,
        [
            "Business-Outcomes (z. B. OEE, Kundenerlebnis) mit technischer Basis (MQTT) verbinden.",
            "Standard-Compliance und Vendor-Neutralität betonen.",
        ],
    )

    doc.add_page_break()

    # Sprecher:in 2 – Produkt & Technologie
    add_heading(doc, "Sprecher:in 2 – Produkt & Technologie", level=1)

    add_heading(doc, "Folie 5 – MQTT Basics in 60 Sekunden", level=2)
    add_bullets(
        doc,
        [
            "Publish/Subscribe, Topics, QoS 0/1/2",
            "Retained Messages, Last Will & Testament",
        ],
    )

    add_heading(doc, "Folie 6 – HiveMQ Plattformüberblick", level=2)
    add_bullets(
        doc,
        [
            "Komponenten: Broker/Cluster, HiveMQ Cloud (Managed), Data Hub, Extensions SDK, Observability",
            "Deployment: Kubernetes/VM, Multi‑Cloud, Edge‑Integration",
        ],
    )

    add_heading(doc, "Folie 7 – Enterprise‑Eigenschaften", level=2)
    add_bullets(
        doc,
        [
            "Zuverlässigkeit: Hohe Verfügbarkeit, Stateful Sessions, Backpressure",
            "Sicherheit: TLS, Authentifizierung/Autorisierung, Zertifikate, Richtlinien",
            "Operability: Metrics, Tracing, Integrationen (Data Lakes, Analytics, iPaaS)",
        ],
    )

    add_heading(doc, "Folie 8 – Use Case: Connected Factory", level=2)
    add_bullets(
        doc,
        [
            "Flow: Sensor → MQTT → HiveMQ → Analytics/ERP",
            "Value: Echtzeit‑Monitoring, geringere Ausfallzeiten, Datenqualität",
        ],
    )
    add_notes(
        doc,
        [
            "Fokus auf Betriebsrelevanz: Latenz, Zuverlässigkeit, Betrieb, nicht zu tief in Protokolldetails.",
            "Optional Mini‑Demo: MQTT-Client zeigt Publish/Subscribe gegen Test‑Broker.",
        ],
    )

    doc.add_page_break()

    # Sprecher:in 3 – Geschäftsmodell, Go-to-Market, Karriere
    add_heading(doc, "Sprecher:in 3 – Geschäftsmodell, Go‑to‑Market, Karriere", level=1)

    add_heading(doc, "Folie 9 – Geschäftsmodell & Pricing‑Hebel", level=2)
    add_bullets(
        doc,
        [
            "Subscription/Enterprise‑Lizenz, Managed Service (Cloud)",
            "Werttreiber: Verbindungsanzahl/Throughput, Cluster‑Größe, SLA, Add‑ons (Data Hub)",
            "TCO: Build‑vs‑Buy – geringere Betriebs‑ und Fehlerrisiken",
        ],
    )

    add_heading(doc, "Folie 10 – Go‑to‑Market & Wettbewerb", level=2)
    add_bullets(
        doc,
        [
            "GTM: Direktvertrieb Enterprise, Partner/Integratoren, Cloud‑Angebot",
            "Wettbewerb: Open Source Broker (Mosquitto, VerneMQ), kommerzielle (EMQX), Hyperscaler‑Dienste",
            "Differenzierung: Enterprise‑Reliability, Operability, Data Quality, Support",
        ],
    )

    add_heading(doc, "Folie 11 – Karrierepfade für Wirtschaftsinformatik", level=2)
    add_bullets(
        doc,
        [
            "Rollen: Product Management, Solutions/Pre‑Sales, Customer Success, Data/Cloud Engineering, Partner Management",
            "Skill‑Fit: Schnittstelle Business‑Tech, Architekturen, Datenintegration, Security/Governance",
        ],
    )

    add_heading(doc, "Folie 12 – Takeaways & Q&A", level=2)
    add_bullets(
        doc,
        [
            "HiveMQ = Brückentechnologie zwischen Geräten und Business‑Systemen",
            "Standardbasiert, skalierbar, betriebssicher – beschleunigt IoT‑Value",
            "Kurze Q&A zum Abschluss",
        ],
    )
    add_notes(
        doc,
        [
            "Klar machen, wie WiInf Mehrwert schafft: Business‑Ziele in Datenflüsse/Policies/SLAs übersetzen.",
        ],
    )

    doc.add_page_break()

    # Visuals, Interaktive Elemente, Weiterführend
    add_heading(doc, "Visual‑Vorschläge", level=1)
    add_bullets(
        doc,
        [
            "Architektur‑Diagramm: Geräte → Broker‑Cluster → Integrationen (Analytics/ERP/Cloud)",
            "Value Map: Technik‑Eigenschaften → Business‑KPIs (Uptime, Time‑to‑Market, OEE)",
            "Wettbewerbsradar: Reliability/Operability vs. Flexibilität/Ökosystem",
        ],
    )

    add_heading(doc, "Interaktive Elemente", level=1)
    add_bullets(
        doc,
        [
            "30‑Sekunden‑Umfrage: Wer hat schon MQTT genutzt?",
            "Mini‑Übung: Use Case in 3 Topics skizzieren (Gerät, Linie, Werk)",
        ],
    )

    add_heading(doc, "Handout & Weiterführendes", level=1)
    add_bullets(
        doc,
        [
            "MQTT Standard (OASIS): Grundlagen und QoS‑Konzepte",
            "HiveMQ Website: https://www.hivemq.com",
            "Best Practices: Architektur & Operations für IoT‑Skalierung",
        ],
    )

    return doc


def main() -> None:
    output_path = "/workspace/HiveMQ_Vortrag_Wirtschaftsinformatik.docx"
    document = build_document()
    # Set base font size for readability
    style = document.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    document.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == "__main__":
    main()

